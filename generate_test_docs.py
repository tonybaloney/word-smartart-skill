"""
generate_test_docs.py - Generate test Word documents demonstrating each SmartArt type.

Run this script to create a set of .docx files in the tests/ directory,
each showcasing a different SmartArt diagram type.

Usage:
    python generate_test_docs.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from smartart import SmartArt


OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests")


def ensure_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def generate_basic_list():
    """Test 1: Basic Block List diagram."""
    doc = Document()
    doc.add_heading("SmartArt Test: Basic Block List", 0)
    doc.add_paragraph(
        "This document tests the Basic Block List SmartArt layout. "
        "Each item should appear as a separate block in a grid arrangement."
    )

    SmartArt.add_basic_list(doc, "Programming Languages", [
        "Python",
        "JavaScript",
        "Rust",
        "Go",
        "TypeScript",
        "C#",
    ])

    doc.add_paragraph("")  # spacer

    SmartArt.add_basic_list(doc, "Project Priorities", [
        "Security",
        "Performance",
        "Usability",
    ])

    path = os.path.join(OUTPUT_DIR, "test_basic_list.docx")
    SmartArt.save(doc, path)
    print(f"  Created: {path}")
    return path


def generate_basic_process():
    """Test 2: Basic Process (linear flow) diagram."""
    doc = Document()
    doc.add_heading("SmartArt Test: Basic Process", 0)
    doc.add_paragraph(
        "This document tests the Basic Process SmartArt layout. "
        "Items should appear as sequential steps with arrows between them."
    )

    SmartArt.add_basic_process(doc, "Software Development Lifecycle", [
        "Requirements",
        "Design",
        "Implementation",
        "Testing",
        "Deployment",
        "Maintenance",
    ])

    doc.add_paragraph("")

    SmartArt.add_basic_process(doc, "CI/CD Pipeline", [
        "Commit",
        "Build",
        "Test",
        "Deploy",
    ])

    path = os.path.join(OUTPUT_DIR, "test_basic_process.docx")
    SmartArt.save(doc, path)
    print(f"  Created: {path}")
    return path


def generate_hierarchy():
    """Test 3: Hierarchy (org chart) diagram."""
    doc = Document()
    doc.add_heading("SmartArt Test: Hierarchy", 0)
    doc.add_paragraph(
        "This document tests the Hierarchy SmartArt layout. "
        "Items should appear in a tree/org chart structure."
    )

    SmartArt.add_hierarchy(doc, "Company Organization", {
        "CEO": {
            "VP Engineering": {
                "Frontend Team Lead": {
                    "Senior Developer": {},
                    "Junior Developer": {},
                },
                "Backend Team Lead": {
                    "Senior Developer": {},
                    "DevOps Engineer": {},
                },
            },
            "VP Marketing": {
                "Brand Manager": {},
                "Content Lead": {},
            },
            "VP Sales": {
                "Enterprise Sales": {},
                "SMB Sales": {},
            },
        }
    })

    doc.add_paragraph("")

    SmartArt.add_hierarchy(doc, "File System Structure", {
        "root/": {
            "src/": {
                "components/": {},
                "utils/": {},
            },
            "tests/": {
                "unit/": {},
                "integration/": {},
            },
            "docs/": {},
        }
    })

    path = os.path.join(OUTPUT_DIR, "test_hierarchy.docx")
    SmartArt.save(doc, path)
    print(f"  Created: {path}")
    return path


def generate_cycle():
    """Test 4: Cycle diagram."""
    doc = Document()
    doc.add_heading("SmartArt Test: Cycle", 0)
    doc.add_paragraph(
        "This document tests the Cycle SmartArt layout. "
        "Items should appear in a circular arrangement suggesting a repeating process."
    )

    SmartArt.add_cycle(doc, "Agile Sprint Cycle", [
        "Sprint Planning",
        "Development",
        "Code Review",
        "Testing",
        "Sprint Review",
        "Retrospective",
    ])

    doc.add_paragraph("")

    SmartArt.add_cycle(doc, "Data Science Workflow", [
        "Collect Data",
        "Clean & Prepare",
        "Explore & Visualize",
        "Model & Train",
        "Evaluate",
        "Deploy & Monitor",
    ])

    path = os.path.join(OUTPUT_DIR, "test_cycle.docx")
    SmartArt.save(doc, path)
    print(f"  Created: {path}")
    return path


def generate_pyramid():
    """Test 5: Pyramid diagram."""
    doc = Document()
    doc.add_heading("SmartArt Test: Pyramid", 0)
    doc.add_paragraph(
        "This document tests the Pyramid SmartArt layout. "
        "Items should appear stacked in a pyramid shape, top to bottom."
    )

    SmartArt.add_pyramid(doc, "Testing Pyramid", [
        "E2E Tests",
        "Integration Tests",
        "Unit Tests",
    ])

    doc.add_paragraph("")

    SmartArt.add_pyramid(doc, "Maslow's Hierarchy of Developer Needs", [
        "Self-Actualization (Open Source Contributions)",
        "Esteem (Conference Talks)",
        "Belonging (Team & Community)",
        "Safety (Job Security & Good Tooling)",
        "Physiological (Coffee & Fast WiFi)",
    ])

    path = os.path.join(OUTPUT_DIR, "test_pyramid.docx")
    SmartArt.save(doc, path)
    print(f"  Created: {path}")
    return path


def generate_radial():
    """Test 6: Radial (hub and spoke) diagram."""
    doc = Document()
    doc.add_heading("SmartArt Test: Radial", 0)
    doc.add_paragraph(
        "This document tests the Radial SmartArt layout. "
        "A central item should appear with surrounding items radiating outward."
    )

    SmartArt.add_radial(doc, "Microservices Architecture", "API Gateway", [
        "Auth Service",
        "User Service",
        "Payment Service",
        "Notification Service",
        "Analytics Service",
    ])

    doc.add_paragraph("")

    SmartArt.add_radial(doc, "Core Skills for Developers", "Problem Solving", [
        "Communication",
        "Technical Writing",
        "Code Review",
        "Debugging",
        "Architecture",
    ])

    path = os.path.join(OUTPUT_DIR, "test_radial.docx")
    SmartArt.save(doc, path)
    print(f"  Created: {path}")
    return path


def generate_combined():
    """Test 7: Combined document with multiple SmartArt types."""
    doc = Document()
    doc.add_heading("SmartArt Showcase: All Diagram Types", 0)
    doc.add_paragraph(
        "This document demonstrates all supported SmartArt diagram types "
        "in a single document, proving that multiple diagrams can coexist."
    )

    # 1. Process
    SmartArt.add_basic_process(doc, "1. Process: How a Feature Gets Built", [
        "Idea",
        "RFC",
        "Prototype",
        "Review",
        "Ship",
    ])

    doc.add_paragraph("")

    # 2. List
    SmartArt.add_basic_list(doc, "2. List: Tech Stack", [
        "React",
        "Node.js",
        "PostgreSQL",
        "Redis",
    ])

    doc.add_paragraph("")

    # 3. Hierarchy
    SmartArt.add_hierarchy(doc, "3. Hierarchy: Module Structure", {
        "Application": {
            "Frontend": {
                "Components": {},
                "Pages": {},
            },
            "Backend": {
                "API Routes": {},
                "Database": {},
            },
        }
    })

    doc.add_paragraph("")

    # 4. Cycle
    SmartArt.add_cycle(doc, "4. Cycle: DevOps Loop", [
        "Plan",
        "Code",
        "Build",
        "Test",
        "Release",
        "Monitor",
    ])

    doc.add_paragraph("")

    # 5. Pyramid
    SmartArt.add_pyramid(doc, "5. Pyramid: Priority Levels", [
        "Critical",
        "High",
        "Medium",
        "Low",
    ])

    doc.add_paragraph("")

    # 6. Radial
    SmartArt.add_radial(doc, "6. Radial: System Components", "Core Engine", [
        "Plugin System",
        "Config Manager",
        "Event Bus",
        "Logger",
    ])

    path = os.path.join(OUTPUT_DIR, "test_combined_all_types.docx")
    SmartArt.save(doc, path)
    print(f"  Created: {path}")
    return path


def main():
    ensure_output_dir()
    print("Generating SmartArt test documents...")
    print()

    paths = []
    paths.append(generate_basic_list())
    paths.append(generate_basic_process())
    paths.append(generate_hierarchy())
    paths.append(generate_cycle())
    paths.append(generate_pyramid())
    paths.append(generate_radial())
    paths.append(generate_combined())

    print()
    print(f"Generated {len(paths)} test documents in {OUTPUT_DIR}/")
    print()

    # Verify each file is a valid ZIP (basic .docx validation)
    import zipfile
    all_valid = True
    for path in paths:
        if zipfile.is_zipfile(path):
            with zipfile.ZipFile(path, 'r') as z:
                has_diagrams = any('diagrams/' in n for n in z.namelist())
                status = "OK (has diagrams)" if has_diagrams else "WARN (no diagrams)"
        else:
            status = "FAIL (not a valid ZIP)"
            all_valid = False
        print(f"  {os.path.basename(path)}: {status}")

    print()
    if all_valid:
        print("All documents generated and validated successfully!")
    else:
        print("Some documents failed validation.")
        exit(1)


if __name__ == "__main__":
    main()
