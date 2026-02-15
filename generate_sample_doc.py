"""
generate_sample_doc.py - Example of combining Markdown content with SmartArt diagrams.

This demonstrates the workflow an AI agent would follow when asked:
"Generate a Word doc explaining a complicated topic and include diagrams."

The agent writes content as Markdown, converts sections to Word via python-docx,
and inserts SmartArt diagrams at appropriate points using the smartart skill.

Usage:
    python generate_sample_doc.py
"""

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from smartart import SmartArt


def add_markdown_to_doc(doc: Document, md_text: str):
    """
    Convert a block of Markdown text to Word paragraphs and add to the document.

    Supports: headings (h1-h4), paragraphs, bold, italic, bullet lists, numbered lists.
    This is the pattern an AI agent would use to write prose content.
    """
    html = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")

    for element in soup.children:
        if element.name and element.name.startswith("h"):
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)
        elif element.name == "p":
            _add_rich_paragraph(doc, element)
        elif element.name in ("ul", "ol"):
            style = "List Bullet" if element.name == "ul" else "List Number"
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style=style)
        elif element.name == "pre":
            code = element.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif element.name == "table":
            _add_table(doc, element)


def _add_rich_paragraph(doc, element):
    """Add a paragraph with bold/italic inline formatting."""
    p = doc.add_paragraph()
    for child in element.children:
        if isinstance(child, str):
            p.add_run(child)
        elif child.name == "strong":
            run = p.add_run(child.get_text())
            run.bold = True
        elif child.name == "em":
            run = p.add_run(child.get_text())
            run.italic = True
        elif child.name == "code":
            run = p.add_run(child.get_text())
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            p.add_run(child.get_text())


def _add_table(doc, table_element):
    """Add an HTML table to the document."""
    rows_data = []
    for tr in table_element.find_all("tr"):
        cells = [td.get_text().strip() for td in tr.find_all(["th", "td"])]
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(row) for row in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols, style="Light Grid Accent 1")
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                table.rows[i].cells[j].text = cell_text


def generate_kubernetes_explainer():
    """
    Generate a comprehensive Word document explaining Kubernetes architecture
    with SmartArt diagrams illustrating key concepts.
    """
    doc = Document()

    # ── Title and Introduction ──────────────────────────────────────────

    add_markdown_to_doc(doc, """
# Understanding Kubernetes Architecture

Kubernetes (K8s) is an open-source container orchestration platform that automates
the deployment, scaling, and management of containerized applications. Originally
designed by Google and now maintained by the Cloud Native Computing Foundation (CNCF),
Kubernetes has become the industry standard for running distributed systems.

This document explains the core architecture, components, and workflows that make
Kubernetes work.
""")

    # ── High-Level Architecture Diagram ─────────────────────────────────

    SmartArt.add_hierarchy(doc, "Kubernetes Architecture Overview", {
        "Kubernetes Cluster": {
            "Control Plane": {
                "API Server": {},
                "etcd": {},
                "Scheduler": {},
                "Controller Manager": {},
            },
            "Worker Nodes": {
                "Node 1": {},
                "Node 2": {},
                "Node N": {},
            },
        }
    })

    # ── Control Plane Section ───────────────────────────────────────────

    add_markdown_to_doc(doc, """
## The Control Plane

The control plane is the brain of the Kubernetes cluster. It makes global decisions
about the cluster (such as scheduling) and detects and responds to cluster events
(such as starting up a new pod when a deployment's replica count is unsatisfied).

The control plane components can run on any machine in the cluster, but for simplicity
they are typically all started on the same machine, separate from user containers.
""")

    SmartArt.add_basic_list(doc, "Control Plane Components", [
        "API Server — Front door to the cluster. All communication goes through it.",
        "etcd — Distributed key-value store. The cluster's single source of truth.",
        "Scheduler — Assigns newly created pods to nodes based on resource requirements.",
        "Controller Manager — Runs controller loops that regulate the state of the cluster.",
    ])

    add_markdown_to_doc(doc, """
### How the Control Plane Processes a Request

When you run `kubectl apply -f deployment.yaml`, a precise sequence of events occurs
within the control plane to make your desired state a reality.
""")

    SmartArt.add_basic_process(doc, "Request Processing Flow", [
        "kubectl sends request to API Server",
        "API Server validates and persists to etcd",
        "Scheduler detects unscheduled pods",
        "Scheduler assigns pods to nodes",
        "Kubelet on node pulls container images",
        "Containers start running",
    ])

    # ── Worker Nodes Section ────────────────────────────────────────────

    add_markdown_to_doc(doc, """
## Worker Nodes

Worker nodes are the machines that run your containerized applications. Every worker
node runs at minimum:

- **Kubelet** — An agent that ensures containers are running in a pod
- **Container Runtime** — The software responsible for running containers (e.g., containerd, CRI-O)
- **Kube-proxy** — A network proxy that maintains network rules on nodes
""")

    SmartArt.add_radial(doc, "Worker Node Components", "Kubelet", [
        "Container Runtime",
        "Kube-proxy",
        "Pods",
        "cAdvisor (Monitoring)",
    ])

    # ── Pod Lifecycle Section ───────────────────────────────────────────

    add_markdown_to_doc(doc, """
## Pod Lifecycle

A Pod is the smallest deployable unit in Kubernetes. It represents a single instance
of a running process and can contain one or more containers. Pods go through a
well-defined lifecycle from creation to termination.

Understanding the pod lifecycle is critical for debugging application issues and
designing resilient systems.
""")

    SmartArt.add_cycle(doc, "Pod Lifecycle States", [
        "Pending — Accepted but not yet running",
        "Running — At least one container is running",
        "Succeeded — All containers terminated successfully",
        "Failed — At least one container terminated in failure",
        "Unknown — State cannot be determined",
    ])

    # ── Deployment Strategy Section ─────────────────────────────────────

    add_markdown_to_doc(doc, """
## Deployment Strategies

Kubernetes supports multiple strategies for updating your applications with zero
or minimal downtime. Choosing the right strategy depends on your application's
tolerance for downtime, risk appetite, and infrastructure constraints.

| Strategy | Downtime | Risk | Rollback Speed |
|----------|----------|------|----------------|
| Rolling Update | None | Low | Fast |
| Recreate | Yes | Low | Fast |
| Blue/Green | None | Medium | Instant |
| Canary | None | Low | Fast |
""")

    SmartArt.add_basic_process(doc, "Rolling Update Process", [
        "New ReplicaSet created",
        "Scale up new pods",
        "Health checks pass",
        "Scale down old pods",
        "Repeat until complete",
    ])

    # ── Networking Section ──────────────────────────────────────────────

    add_markdown_to_doc(doc, """
## Kubernetes Networking Model

Kubernetes networking addresses four concerns:

1. **Container-to-container** communications within a Pod (via localhost)
2. **Pod-to-pod** communications across nodes (flat network)
3. **Pod-to-service** communications (via virtual IPs)
4. **External-to-service** communications (via Ingress or LoadBalancer)

The networking model requires that all pods can communicate with each other without
NAT, creating a flat network space across the entire cluster.
""")

    SmartArt.add_pyramid(doc, "Kubernetes Networking Layers", [
        "Ingress / Load Balancer — External traffic entry",
        "Services — Stable endpoints and load balancing",
        "Pod Network — Flat network between all pods",
        "Container Network — Localhost within a pod",
    ])

    # ── Summary Section ─────────────────────────────────────────────────

    add_markdown_to_doc(doc, """
## Key Takeaways

Kubernetes provides a powerful, declarative platform for running containerized
applications at scale. Its architecture separates concerns between the **control plane**
(which makes decisions) and **worker nodes** (which execute work), creating a
resilient and scalable system.

The key to working effectively with Kubernetes is understanding that it operates on
a **desired state** model — you tell Kubernetes what you want, and its controllers
continuously work to make reality match your declaration.
""")

    SmartArt.add_basic_list(doc, "Summary: Core Kubernetes Concepts", [
        "Declarative Configuration",
        "Self-Healing Systems",
        "Horizontal Scaling",
        "Service Discovery",
        "Rolling Updates",
        "Resource Management",
    ])

    # Save
    output_path = "tests/sample_kubernetes_explainer.docx"
    doc.save(output_path)

    # Finalize: open in Word to regenerate SmartArt rendering
    SmartArt.finalize(output_path)

    print("Generated: {}".format(output_path))
    return output_path


if __name__ == "__main__":
    path = generate_kubernetes_explainer()

    # Validate
    import zipfile
    with zipfile.ZipFile(path, 'r') as z:
        diagrams = [n for n in z.namelist() if 'diagrams/' in n]
        print("  Diagram parts: {} (from {} SmartArt objects)".format(
            len(diagrams), len(diagrams) // 5
        ))
    print("  Open in Microsoft Word to verify rendering.")
