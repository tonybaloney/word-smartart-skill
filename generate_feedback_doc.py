"""
Generate a Word document explaining to the Word team why SmartArt was so hard
to produce programmatically, using the SmartArt skill itself for diagrams.

Output: _feedback_for_word_team.docx (underscore prefix = not committed)
"""
import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from smartart import SmartArt


def add_body(doc, text):
    """Add a body paragraph."""
    doc.add_paragraph(text)


def add_body_bold_lead(doc, bold_text, rest_text):
    """Add a paragraph with bold lead-in text."""
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    p.add_run(rest_text)


def build_document():
    doc = Document()

    # Title
    title = doc.add_heading("Why It Was So Hard to Programmatically Create SmartArt", level=0)

    doc.add_paragraph(
        "This document was generated entirely in Python using the word-smartart-skill library. "
        "Every diagram below is a native, editable SmartArt object\u2014not an image. "
        "The fact that you can click on any diagram and edit it in Word is the whole point of this project, "
        "and also what made it so difficult to build."
    )

    # ── Section 1: The Goal ──
    doc.add_heading("The Goal", level=1)
    add_body(doc,
        "Build a GitHub Copilot Skill that teaches AI coding agents how to insert native SmartArt "
        "diagrams into .docx files. When an AI generates a Word document, it should include real, "
        "editable SmartArt\u2014not rasterized images of diagrams."
    )

    SmartArt.add_basic_list(doc, "Requirements", [
        "No COM automation at runtime",
        "No inserted images",
        "Valid OpenXML only",
        "Renders in Word without repair",
        "Diagrams are editable",
    ])

    # ── Section 2: What a SmartArt Actually Is ──
    doc.add_heading("What a Single SmartArt Diagram Actually Requires", level=1)
    add_body(doc,
        "A single SmartArt diagram is not one XML element. It is five separate XML parts "
        "inside the .docx ZIP archive, each with its own content type and relationship URI. "
        "Getting any one of these wrong results in silent failure."
    )

    SmartArt.add_basic_list(doc, "The Five Required XML Parts", [
        "data.xml \u2013 Data model with nodes and connections",
        "layout.xml \u2013 Layout algorithm definition",
        "quickStyle.xml \u2013 Visual style definitions",
        "colors.xml \u2013 Color scheme",
        "drawing.xml \u2013 Cached fallback rendering",
    ])

    # ── Section 3: The Undocumented Presentation Points ──
    doc.add_heading("The Undocumented Presentation Points", level=1)
    add_body(doc,
        "This was the single hardest problem. The ECMA-376 spec and every blog post suggest "
        "that a data model needs: a doc point, data points, transition points, and parOf connections. "
        "This structure is valid XML. But Word renders it as a collapsed zero-height rectangle."
    )
    add_body(doc,
        "The missing piece: Word also requires presentation points (type=\"pres\") and presentation "
        "connections (type=\"presOf\", type=\"presParOf\"). For a simple 8-item list:"
    )

    SmartArt.add_hierarchy(doc, "What the Spec Describes vs. What Word Needs", {
        "Data Model": {
            "Spec says you need": {
                "1 doc point": {},
                "N data points": {},
                "N parTrans points": {},
                "N sibTrans points": {},
                "N parOf connections": {},
            },
            "Word also requires": {
                "~2N pres points": {},
                "~N presOf connections": {},
                "~2N presParOf connections": {},
            },
        }
    })

    add_body(doc,
        "Without the presentation points, Word\u2019s rendering engine has nothing to work with. "
        "There is no documentation explaining how many are needed, what attributes they require, "
        "or how they vary by layout type. The only way to discover this was to create SmartArt "
        "in Word via the GUI, save, unzip, and reverse-engineer the XML."
    )

    # ── Section 4: Relationship Type Minefield ──
    doc.add_heading("The Relationship Type Minefield", level=1)
    add_body(doc,
        "Three of the five SmartArt parts use a microsoft.com/office/2007 namespace instead of "
        "the standard openxmlformats.org namespace. If you use the wrong (but seemingly correct) "
        "namespace, Word silently ignores the parts. No error. No warning. Just a blank space."
    )

    SmartArt.add_basic_list(doc, "Parts Using Non-Standard Namespaces", [
        "Style \u2013 schemas.microsoft.com/office/2007 (NOT openxmlformats.org)",
        "Colors \u2013 schemas.microsoft.com/office/2007 (NOT openxmlformats.org)",
        "Drawing \u2013 schemas.microsoft.com/office/2007 (NOT openxmlformats.org)",
    ])

    # ── Section 5: Silent Failures ──
    doc.add_heading("Silent Failures: The Debugging Nightmare", level=1)
    add_body(doc,
        "When SmartArt XML is malformed, Word does not show an error, log a warning, display "
        "a placeholder, offer to repair, or indicate which part is broken. It silently renders "
        "a zero-height invisible rectangle. This led to an excruciating debugging cycle:"
    )

    SmartArt.add_cycle(doc, "The SmartArt Debugging Cycle", [
        "Modify XML",
        "Re-zip as .docx",
        "Open in Word",
        "See nothing rendered",
        "Diff against working file",
        "Guess what\u2019s wrong",
    ])

    # ── Section 6: Our 8 Attempts ──
    doc.add_heading("Eight Attempts to Insert a Simple List", level=1)
    add_body(doc,
        "The following shows the progression of approaches we tried before finding one that works. "
        "Each attempt required significant development time and testing."
    )

    SmartArt.add_basic_process(doc, "Attempt Timeline", [
        "Generate all XML from scratch \u2013 FAILED",
        "Fix relationship URIs \u2013 FAILED",
        "Add placeholder pres points \u2013 FAILED",
        "Extract templates via COM \u2013 FAILED",
        "Fix compatibility mode \u2013 FAILED",
        "Copy template drawing XML \u2013 FAILED",
        "COM finalize post-process \u2013 FAILED",
        "Modify template data in-place \u2013 SUCCESS",
    ])

    # ── Section 7: The Solution ──
    doc.add_heading("The Solution: Template-Based Generation", level=1)
    add_body(doc,
        "After seven failed approaches, we arrived at the only reliable method: use Word COM "
        "automation as a one-time step to create template documents, then modify the template\u2019s "
        "data model XML in-place at generation time, preserving the presentation points."
    )

    SmartArt.add_basic_process(doc, "How It Works", [
        "One-time: Generate templates via Word COM",
        "Extract 5 XML parts from template",
        "Modify data model text in-place",
        "Preserve presentation points",
        "Inject into target .docx",
    ])

    # ── Section 8: What Would Help ──
    doc.add_heading("What Would Make This Feasible", level=1)
    add_body(doc,
        "If the Word team wanted to make SmartArt accessible to AI agents and programmatic tools, "
        "these changes would have the highest impact:"
    )

    SmartArt.add_pyramid(doc, "Recommendations by Impact (Top = Highest)", [
        "Support data-only mode (auto-generate pres points on open)",
        "Add SmartArt to Open XML SDK",
        "Provide a CLI tool for SmartArt XML generation",
        "Document the presentation point requirements",
        "Add error reporting for malformed SmartArt",
        "Document the exact relationship URIs",
    ])

    # ── Section 9: The Central Problem ──
    doc.add_heading("The Core Issue", level=1)
    add_body(doc,
        "SmartArt is a genuinely useful feature for end users. But from a developer\u2019s perspective "
        "\u2014and especially from an AI agent\u2019s perspective\u2014it is a black box wrapped in "
        "undocumented XML, protected by silent failures, and accessible only through reverse engineering."
    )

    SmartArt.add_radial(doc, "SmartArt's Accessibility Problem", "SmartArt Feature", [
        "Undocumented pres points",
        "Silent zero-height failures",
        "Inconsistent namespace URIs",
        "Opaque layout definitions",
        "No SDK support anywhere",
        "Compatibility mode gotchas",
    ])

    add_body(doc,
        "As AI-generated documents become more common, the gap between \u201cwhat Word can render\u201d "
        "and \u201cwhat tools can generate\u201d will become a real limitation. The Word team has an "
        "opportunity to make SmartArt the showcase example of how Office formats can be AI-friendly. "
        "The layout engine is already doing the hard work\u2014it just needs to be accessible."
    )

    # ── Footer note ──
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(
        "This document was generated programmatically using python-docx and the word-smartart-skill library. "
        "Every diagram is a native SmartArt object. Right-click any diagram to verify."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


if __name__ == "__main__":
    doc = build_document()
    out_path = os.path.join(os.path.dirname(__file__), "_feedback_for_word_team.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")

    # Verify diagram parts exist
    import zipfile
    with zipfile.ZipFile(out_path, 'r') as z:
        diagram_parts = [n for n in z.namelist() if 'diagrams/' in n]
        print(f"Diagram parts in document: {len(diagram_parts)}")
        data_parts = [n for n in diagram_parts if 'data' in n]
        print(f"SmartArt diagrams: {len(data_parts)}")
