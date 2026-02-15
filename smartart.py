"""
smartart.py - Create native SmartArt (DrawingML Diagrams) in Word documents.

This module creates real, editable SmartArt objects in .docx files. It works by
extracting diagram XML parts from Word-generated templates and injecting modified
versions into target documents using python-docx and lxml.

No images are generated - these are native SmartArt objects that render and are
fully editable in Microsoft Word.

Supported diagram types:
- Basic List: Simple block list of items
- Basic Process: Linear flow from left to right
- Hierarchy: Org chart / tree structure
- Cycle: Circular process
- Pyramid: Layered pyramid
- Radial: Central item with radiating connections

Requirements:
    pip install python-docx lxml

Usage:
    from docx import Document
    from smartart import SmartArt

    doc = Document()
    doc.add_heading("My Document", 0)

    SmartArt.add_basic_process(doc, "My Process", [
        "Step 1: Research",
        "Step 2: Design",
        "Step 3: Implement",
        "Step 4: Test",
    ])

    doc.save("output.docx")
"""

import os
import copy
import uuid
import zipfile
from typing import Optional
from lxml import etree
from docx import Document
from docx.opc.part import Part
from docx.opc.packuri import PackURI

# ─── Paths ──────────────────────────────────────────────────────────────────

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")

# ─── OpenXML Namespaces ────────────────────────────────────────────────────

NS = {
    "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "dsp": "http://schemas.microsoft.com/office/drawing/2008/diagram",
}

# Content types for SmartArt parts
CT_DGM_DATA = "application/vnd.openxmlformats-officedocument.drawingml.diagramData+xml"
CT_DGM_LAYOUT = "application/vnd.openxmlformats-officedocument.drawingml.diagramLayout+xml"
CT_DGM_STYLE = "application/vnd.openxmlformats-officedocument.drawingml.diagramStyle+xml"
CT_DGM_COLORS = "application/vnd.openxmlformats-officedocument.drawingml.diagramColors+xml"
CT_DGM_DRAWING = "application/vnd.ms-office.drawingml.diagramDrawing+xml"

# Relationship types for diagram parts (must match what Word expects)
RT_DGM_DATA = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramData"
RT_DGM_LAYOUT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramLayout"
RT_DGM_STYLE = "http://schemas.microsoft.com/office/2007/relationships/diagramStyle"
RT_DGM_COLORS = "http://schemas.microsoft.com/office/2007/relationships/diagramColors"
RT_DGM_DRAWING = "http://schemas.microsoft.com/office/2007/relationships/diagramDrawing"

# Map template names to layout URIs
LAYOUT_URIS = {
    "basic_list": "urn:microsoft.com/office/officeart/2005/8/layout/default",
    "basic_process": "urn:microsoft.com/office/officeart/2005/8/layout/process1",
    "hierarchy": "urn:microsoft.com/office/officeart/2005/8/layout/hierarchy1",
    "cycle": "urn:microsoft.com/office/officeart/2005/8/layout/cycle1",
    "pyramid": "urn:microsoft.com/office/officeart/2005/8/layout/pyramid1",
    "radial": "urn:microsoft.com/office/officeart/2005/8/layout/radial1",
}


def _qn(ns: str, tag: str) -> str:
    """Create a qualified name {namespace}tag."""
    return "{{{}}}{}".format(NS[ns], tag)


def _new_guid() -> str:
    """Generate a GUID in SmartArt format."""
    return "{{{}}}".format(str(uuid.uuid4()).upper())


def _ensure_modern_compat(doc: Document):
    """
    Upgrade the document's compatibility mode to Word 2013+ (val=15).

    python-docx defaults to compatibilityMode=14 (Word 2010), which causes
    Word to open in "Compatibility Mode" and not regenerate SmartArt diagrams.
    Setting it to 15 ensures modern rendering.
    """
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    settings_part = doc.settings.element

    compat = settings_part.find("{%s}compat" % W_NS)
    if compat is None:
        compat = etree.SubElement(settings_part, "{%s}compat" % W_NS)

    # Find or create the compatibilityMode setting
    for setting in compat.findall("{%s}compatSetting" % W_NS):
        if setting.get("{%s}name" % W_NS) == "compatibilityMode":
            setting.set("{%s}val" % W_NS, "15")
            return

    # Not found — add it
    etree.SubElement(compat, "{%s}compatSetting" % W_NS, attrib={
        "{%s}name" % W_NS: "compatibilityMode",
        "{%s}uri" % W_NS: "http://schemas.microsoft.com/office/word",
        "{%s}val" % W_NS: "15",
    })


def _next_diagram_id(doc: Document) -> int:
    """Find the next available diagram part index in the document."""
    existing = set()
    for rel in doc.part.rels.values():
        target = rel.target_ref if hasattr(rel, 'target_ref') else ""
        if "diagrams/data" in target:
            try:
                num = int(target.split("data")[1].split(".")[0])
                existing.add(num)
            except (ValueError, IndexError):
                pass
    return max(existing, default=0) + 1


# ─── Template Extraction ───────────────────────────────────────────────────

def _extract_template_parts(template_name: str) -> dict:
    """
    Extract diagram XML parts from a pre-generated Word template.

    Returns dict with keys: 'data', 'layout', 'style', 'colors', 'drawing'
    Each value is the raw XML bytes for that part.
    """
    template_path = os.path.join(TEMPLATE_DIR, "{}.docx".format(template_name))
    if not os.path.exists(template_path):
        raise FileNotFoundError(
            "Template '{}' not found at {}. Run generate_templates.py first.".format(
                template_name, template_path
            )
        )

    parts = {}
    # Map of file name patterns to part keys
    part_map = {
        "data": "data",
        "layout": "layout",
        "quickStyle": "style",
        "colors": "colors",
        "drawing": "drawing",
    }

    with zipfile.ZipFile(template_path, 'r') as z:
        for name in z.namelist():
            if "diagrams/" not in name:
                continue
            for pattern, key in part_map.items():
                if pattern in name:
                    parts[key] = z.read(name)
                    break

    required = {'data', 'layout', 'style', 'colors', 'drawing'}
    missing = required - set(parts.keys())
    if missing:
        raise ValueError("Template missing parts: {}".format(missing))

    return parts


def _extract_rels_from_template(template_name: str) -> dict:
    """Extract the relationship types from the template's document.xml.rels."""
    template_path = os.path.join(TEMPLATE_DIR, "{}.docx".format(template_name))
    rel_types = {}

    with zipfile.ZipFile(template_path, 'r') as z:
        rels_xml = z.read("word/_rels/document.xml.rels")
        root = etree.fromstring(rels_xml)
        for rel in root:
            rtype = rel.get("Type", "")
            target = rel.get("Target", "")
            if "diagramData" in rtype:
                rel_types["data"] = rtype
            elif "diagramLayout" in rtype:
                rel_types["layout"] = rtype
            elif "diagramStyle" in rtype or "diagramQuickStyle" in rtype:
                rel_types["style"] = rtype
            elif "diagramColors" in rtype:
                rel_types["colors"] = rtype
            elif "diagramDrawing" in rtype:
                rel_types["drawing"] = rtype

    return rel_types


# ─── Data Model Modification ───────────────────────────────────────────────

def _modify_template_data_flat(template_data_xml: bytes, items: list[str]) -> bytes:
    """
    Modify a template's data model XML in-place for flat diagrams.

    Instead of building from scratch, this preserves the template's presentation
    points and connections (which are required for rendering) and only changes
    the text content of data nodes. Nodes are added or removed to match the
    requested item count.
    """
    root = etree.fromstring(template_data_xml)
    pt_lst = root.find(_qn("dgm", "ptLst"))
    cxn_lst = root.find(_qn("dgm", "cxnLst"))

    # Find existing data nodes (type=None means data node)
    data_nodes = []
    for pt in pt_lst.findall(_qn("dgm", "pt")):
        if pt.get("type") is None:
            data_nodes.append(pt)

    # Update text of existing nodes
    for i, item_text in enumerate(items):
        if i < len(data_nodes):
            _set_node_text(data_nodes[i], item_text)

    # If we need fewer nodes than the template has, remove extras
    if len(items) < len(data_nodes):
        _remove_excess_nodes(root, pt_lst, cxn_lst, data_nodes, len(items))

    # If we need more nodes than the template has, duplicate the last one
    if len(items) > len(data_nodes):
        _add_extra_nodes(root, pt_lst, cxn_lst, data_nodes, items[len(data_nodes):])

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _set_node_text(pt, text: str):
    """Set the text content of a diagram data point."""
    A_NS = NS["a"]
    t_elem = pt.find(".//{%s}t" % A_NS)
    if t_elem is not None:
        t_elem.text = text
    else:
        # Build text structure if missing
        dgm_t = pt.find(_qn("dgm", "t"))
        if dgm_t is None:
            dgm_t = etree.SubElement(pt, _qn("dgm", "t"))
            etree.SubElement(dgm_t, _qn("a", "bodyPr"))
            etree.SubElement(dgm_t, _qn("a", "lstStyle"))
        # Remove existing paragraphs
        for p in dgm_t.findall(_qn("a", "p")):
            dgm_t.remove(p)
        p = etree.SubElement(dgm_t, _qn("a", "p"))
        r = etree.SubElement(p, _qn("a", "r"))
        etree.SubElement(r, _qn("a", "rPr"), attrib={"lang": "en-US"})
        t = etree.SubElement(r, _qn("a", "t"))
        t.text = text


def _remove_excess_nodes(root, pt_lst, cxn_lst, data_nodes, keep_count: int):
    """Remove data nodes (and their transitions/connections/pres points) beyond keep_count."""
    to_remove_ids = set()
    for node in data_nodes[keep_count:]:
        to_remove_ids.add(node.get("modelId"))

    # Find associated parTrans and sibTrans IDs via connections
    trans_ids = set()
    cxns_to_remove = []
    for cxn in cxn_lst.findall(_qn("dgm", "cxn")):
        dest_id = cxn.get("destId", "")
        src_id = cxn.get("srcId", "")
        if dest_id in to_remove_ids or src_id in to_remove_ids:
            cxns_to_remove.append(cxn)
            par_trans = cxn.get("parTransId", "")
            sib_trans = cxn.get("sibTransId", "")
            if par_trans:
                trans_ids.add(par_trans)
            if sib_trans:
                trans_ids.add(sib_trans)

    all_remove_ids = to_remove_ids | trans_ids

    # Also find presOf connections referencing removed nodes
    pres_ids_to_remove = set()
    for cxn in cxn_lst.findall(_qn("dgm", "cxn")):
        if cxn.get("type") == "presOf" and cxn.get("srcId", "") in all_remove_ids:
            pres_ids_to_remove.add(cxn.get("destId", ""))
            cxns_to_remove.append(cxn)
        elif cxn.get("type") == "presParOf" and cxn.get("destId", "") in pres_ids_to_remove:
            cxns_to_remove.append(cxn)

    all_remove_ids |= pres_ids_to_remove

    # Remove points
    for pt in list(pt_lst.findall(_qn("dgm", "pt"))):
        if pt.get("modelId") in all_remove_ids:
            pt_lst.remove(pt)

    # Remove connections
    for cxn in cxns_to_remove:
        try:
            cxn_lst.remove(cxn)
        except ValueError:
            pass


def _add_extra_nodes(root, pt_lst, cxn_lst, existing_data_nodes, extra_items: list[str]):
    """Add additional data nodes by cloning the last existing node's structure."""
    if not existing_data_nodes:
        return

    last_node = existing_data_nodes[-1]
    last_id = last_node.get("modelId")

    # Find the doc point (parent)
    doc_pt = None
    for pt in pt_lst.findall(_qn("dgm", "pt")):
        if pt.get("type") == "doc":
            doc_pt = pt
            break
    doc_id = doc_pt.get("modelId") if doc_pt is not None else ""

    # Find the parOf connection for the last node to get srcOrd
    last_ord = 0
    for cxn in cxn_lst.findall(_qn("dgm", "cxn")):
        if cxn.get("type") == "parOf" and cxn.get("destId") == last_id:
            last_ord = int(cxn.get("srcOrd", "0"))
            break

    for i, text in enumerate(extra_items):
        new_id = _new_guid()
        par_trans_id = _new_guid()
        sib_trans_id = _new_guid()
        cxn_id = _new_guid()
        ord_val = str(last_ord + i + 1)

        # Data point
        pt = copy.deepcopy(last_node)
        pt.set("modelId", new_id)
        _set_node_text(pt, text)
        pt_lst.append(pt)

        # parTrans point
        par_pt = etree.SubElement(pt_lst, _qn("dgm", "pt"), attrib={
            "modelId": par_trans_id, "type": "parTrans", "cxnId": cxn_id,
        })
        etree.SubElement(par_pt, _qn("dgm", "prSet"))
        etree.SubElement(par_pt, _qn("dgm", "spPr"))

        # sibTrans point
        sib_pt = etree.SubElement(pt_lst, _qn("dgm", "pt"), attrib={
            "modelId": sib_trans_id, "type": "sibTrans", "cxnId": cxn_id,
        })
        etree.SubElement(sib_pt, _qn("dgm", "prSet"))
        etree.SubElement(sib_pt, _qn("dgm", "spPr"))

        # parOf connection
        etree.SubElement(cxn_lst, _qn("dgm", "cxn"), attrib={
            "modelId": cxn_id, "type": "parOf",
            "srcId": doc_id, "destId": new_id,
            "srcOrd": ord_val, "destOrd": "0",
            "parTransId": par_trans_id, "sibTransId": sib_trans_id,
        })


def _modify_template_data_hierarchy(template_data_xml: bytes, hierarchy: dict) -> bytes:
    """
    Modify a hierarchy template's data model with custom hierarchy data.

    Since hierarchy structures vary significantly, we replace data node text
    in a depth-first traversal matching the template's tree structure.
    For now, this modifies existing nodes' text to match the hierarchy.
    """
    root = etree.fromstring(template_data_xml)
    pt_lst = root.find(_qn("dgm", "ptLst"))
    cxn_lst = root.find(_qn("dgm", "cxnLst"))

    # Build a tree structure from the template connections
    doc_pt = None
    nodes_by_id = {}
    for pt in pt_lst.findall(_qn("dgm", "pt")):
        mid = pt.get("modelId")
        ptype = pt.get("type")
        if ptype == "doc":
            doc_pt = pt
        elif ptype is None:
            nodes_by_id[mid] = pt

    # Build parent->children map from parOf connections
    children_map = {}  # parent_id -> [(srcOrd, child_id)]
    for cxn in cxn_lst.findall(_qn("dgm", "cxn")):
        if cxn.get("type") == "parOf":
            src = cxn.get("srcId")
            dst = cxn.get("destId")
            if dst in nodes_by_id:
                children_map.setdefault(src, []).append(
                    (int(cxn.get("srcOrd", "0")), dst)
                )

    # Sort children by srcOrd
    for k in children_map:
        children_map[k].sort()

    # Flatten hierarchy to list for simple text replacement
    def flatten_hierarchy(h):
        result = []
        for label, children_dict in h.items():
            result.append(label)
            if children_dict:
                result.extend(flatten_hierarchy(children_dict))
        return result

    flat_labels = flatten_hierarchy(hierarchy)

    # Walk the template tree and assign text
    def walk_and_assign(parent_id, labels_iter):
        children = children_map.get(parent_id, [])
        for _, child_id in children:
            if child_id in nodes_by_id:
                try:
                    text = next(labels_iter)
                    _set_node_text(nodes_by_id[child_id], text)
                except StopIteration:
                    break
                walk_and_assign(child_id, labels_iter)

    doc_id = doc_pt.get("modelId") if doc_pt is not None else ""
    labels_iter = iter(flat_labels)
    walk_and_assign(doc_id, labels_iter)

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


# ─── Minimal Drawing Part ──────────────────────────────────────────────────

def _build_empty_drawing() -> bytes:
    """Build a minimal drawing part. Word regenerates this on first render."""
    root = etree.Element(_qn("dsp", "drawing"), nsmap={
        "dsp": NS["dsp"],
    })
    sp_tree = etree.SubElement(root, _qn("dsp", "spTree"))
    nvGrpSpPr = etree.SubElement(sp_tree, _qn("dsp", "nvGrpSpPr"))
    etree.SubElement(nvGrpSpPr, _qn("dsp", "cNvPr"), attrib={"id": "0", "name": ""})
    etree.SubElement(nvGrpSpPr, _qn("dsp", "cNvGrpSpPr"))
    etree.SubElement(sp_tree, _qn("dsp", "grpSpPr"))
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


# ─── Part Injection ─────────────────────────────────────────────────────────

def _inject_smartart(doc: Document, template_name: str, data_xml: bytes,
                     width_emu: int = 5486400, height_emu: int = 3200400):
    """
    Inject SmartArt into a document by copying template parts and using custom data.

    Args:
        doc: python-docx Document
        template_name: Name of template (e.g. 'basic_list')
        data_xml: Custom diagram data XML
        width_emu: Width in EMUs
        height_emu: Height in EMUs
    """
    idx = _next_diagram_id(doc)

    # Ensure document is in modern format so Word regenerates SmartArt
    _ensure_modern_compat(doc)

    parts = _extract_template_parts(template_name)
    rel_types = _extract_rels_from_template(template_name)

    # Part names for this diagram
    data_pn = PackURI("/word/diagrams/data{}.xml".format(idx))
    layout_pn = PackURI("/word/diagrams/layout{}.xml".format(idx))
    style_pn = PackURI("/word/diagrams/quickStyle{}.xml".format(idx))
    colors_pn = PackURI("/word/diagrams/colors{}.xml".format(idx))
    drawing_pn = PackURI("/word/diagrams/drawing{}.xml".format(idx))

    # Use custom data, template layout/style/colors, and empty drawing.
    # The drawing part must be empty so Word regenerates it from the data model
    # instead of using a stale cached rendering from the template.
    data_part = Part(data_pn, CT_DGM_DATA, data_xml, doc.part.package)
    layout_part = Part(layout_pn, CT_DGM_LAYOUT, parts['layout'], doc.part.package)
    style_part = Part(style_pn, CT_DGM_STYLE, parts['style'], doc.part.package)
    colors_part = Part(colors_pn, CT_DGM_COLORS, parts['colors'], doc.part.package)
    drawing_part = Part(drawing_pn, CT_DGM_DRAWING, _build_empty_drawing(), doc.part.package)

    # Add relationships using the exact types from the template
    r_data = doc.part.relate_to(data_part, rel_types.get('data', RT_DGM_DATA))
    r_layout = doc.part.relate_to(layout_part, rel_types.get('layout', RT_DGM_LAYOUT))
    r_style = doc.part.relate_to(style_part, rel_types.get('style', RT_DGM_STYLE))
    r_colors = doc.part.relate_to(colors_part, rel_types.get('colors', RT_DGM_COLORS))
    r_drawing = doc.part.relate_to(drawing_part, rel_types.get('drawing', RT_DGM_DRAWING))

    # Build inline drawing paragraph
    para_xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram">'
        '<w:r><w:drawing>'
        '<wp:inline distT="0" distB="0" distL="0" distR="0">'
        '<wp:extent cx="{w}" cy="{h}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:docPr id="{did}" name="Diagram {idx}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/diagram">'
        '<dgm:relIds r:dm="{dm}" r:lo="{lo}" r:qs="{qs}" r:cs="{cs}"/>'
        '</a:graphicData></a:graphic>'
        '</wp:inline>'
        '</w:drawing></w:r></w:p>'
    ).format(
        w=width_emu, h=height_emu, did=idx + 100, idx=idx,
        dm=r_data, lo=r_layout, qs=r_style, cs=r_colors,
    )

    para_element = etree.fromstring(para_xml.encode("utf-8"))
    doc.element.body.append(para_element)


def _add_smartart(doc: Document, template_name: str, items: list[str],
                  title: Optional[str] = None, hierarchy: Optional[dict] = None,
                  width_emu: int = 5486400, height_emu: int = 3200400):
    """
    High-level function to add SmartArt to a document.
    """
    if title:
        doc.add_heading(title, level=2)

    parts = _extract_template_parts(template_name)

    if hierarchy is not None:
        data_xml = _modify_template_data_hierarchy(parts['data'], hierarchy)
    else:
        data_xml = _modify_template_data_flat(parts['data'], items)

    _inject_smartart(doc, template_name, data_xml, width_emu, height_emu)


# ─── Public API ─────────────────────────────────────────────────────────────

class SmartArt:
    """
    Static methods to add SmartArt diagrams to Word documents.

    Each method creates a native, editable SmartArt object - not an image.
    The diagrams will render and be editable in Microsoft Word.

    Templates must be pre-generated (run generate_templates.py once with Word).
    """

    @staticmethod
    def add_basic_list(doc: Document, title: str, items: list[str],
                       width_emu: int = 5486400, height_emu: int = 3200400):
        """
        Add a Basic Block List SmartArt diagram.

        Best for: Ungrouped or non-sequential information, simple lists of items.

        Args:
            doc: python-docx Document object
            title: Heading text above the diagram
            items: List of strings, one per block
            width_emu: Diagram width in EMUs (default ~5.7")
            height_emu: Diagram height in EMUs (default ~3.3")

        Example:
            SmartArt.add_basic_list(doc, "Key Features", [
                "Fast Performance",
                "Easy to Use",
                "Secure by Default",
                "Open Source",
            ])
        """
        _add_smartart(doc, "basic_list", items, title=title,
                      width_emu=width_emu, height_emu=height_emu)

    @staticmethod
    def add_basic_process(doc: Document, title: str, items: list[str],
                          width_emu: int = 5486400, height_emu: int = 3200400):
        """
        Add a Basic Process (linear flow) SmartArt diagram.

        Best for: Sequential steps, workflows, timelines, procedures.

        Args:
            doc: python-docx Document object
            title: Heading text above the diagram
            items: List of strings representing sequential steps
            width_emu: Diagram width in EMUs (default ~5.7")
            height_emu: Diagram height in EMUs (default ~3.3")

        Example:
            SmartArt.add_basic_process(doc, "Development Lifecycle", [
                "Requirements",
                "Design",
                "Implementation",
                "Testing",
                "Deployment",
            ])
        """
        _add_smartart(doc, "basic_process", items, title=title,
                      width_emu=width_emu, height_emu=height_emu)

    @staticmethod
    def add_hierarchy(doc: Document, title: str, hierarchy: dict,
                      width_emu: int = 5486400, height_emu: int = 4000000):
        """
        Add a Hierarchy (org chart / tree) SmartArt diagram.

        Best for: Organizational charts, tree structures, classification.

        Args:
            doc: python-docx Document object
            title: Heading text above the diagram
            hierarchy: Nested dict where keys are labels and values are child dicts.
            width_emu: Diagram width in EMUs (default ~5.7")
            height_emu: Diagram height in EMUs (default ~4.2")

        Example:
            SmartArt.add_hierarchy(doc, "Organization", {
                "CEO": {
                    "VP Engineering": {
                        "Frontend Lead": {},
                        "Backend Lead": {},
                    },
                    "VP Marketing": {
                        "Brand Manager": {},
                    },
                }
            })
        """
        _add_smartart(doc, "hierarchy", [], title=title, hierarchy=hierarchy,
                      width_emu=width_emu, height_emu=height_emu)

    @staticmethod
    def add_cycle(doc: Document, title: str, items: list[str],
                  width_emu: int = 5486400, height_emu: int = 4000000):
        """
        Add a Cycle SmartArt diagram.

        Best for: Repeating processes, continuous workflows, feedback loops.

        Args:
            doc: python-docx Document object
            title: Heading text above the diagram
            items: List of strings representing stages in the cycle
            width_emu: Diagram width in EMUs (default ~5.7")
            height_emu: Diagram height in EMUs (default ~4.2")

        Example:
            SmartArt.add_cycle(doc, "Agile Sprint Cycle", [
                "Plan",
                "Design",
                "Develop",
                "Test",
                "Review",
            ])
        """
        _add_smartart(doc, "cycle", items, title=title,
                      width_emu=width_emu, height_emu=height_emu)

    @staticmethod
    def add_pyramid(doc: Document, title: str, items: list[str],
                    width_emu: int = 5486400, height_emu: int = 4000000):
        """
        Add a Pyramid SmartArt diagram.

        Best for: Proportional or hierarchical relationships, foundation concepts,
        priority levels (top = most important/smallest, bottom = foundation/largest).

        Args:
            doc: python-docx Document object
            title: Heading text above the diagram
            items: List of strings from top (apex) to bottom (base)
            width_emu: Diagram width in EMUs (default ~5.7")
            height_emu: Diagram height in EMUs (default ~4.2")

        Example:
            SmartArt.add_pyramid(doc, "Maslow's Hierarchy", [
                "Self-Actualization",
                "Esteem",
                "Love/Belonging",
                "Safety",
                "Physiological",
            ])
        """
        _add_smartart(doc, "pyramid", items, title=title,
                      width_emu=width_emu, height_emu=height_emu)

    @staticmethod
    def add_radial(doc: Document, title: str, center: str, items: list[str],
                   width_emu: int = 5486400, height_emu: int = 4000000):
        """
        Add a Radial SmartArt diagram with a central item and surrounding items.

        Best for: Relationships to a central concept, components of a system,
        hub-and-spoke architectures.

        Args:
            doc: python-docx Document object
            title: Heading text above the diagram
            center: Text for the central element
            items: List of strings for the surrounding elements
            width_emu: Diagram width in EMUs (default ~5.7")
            height_emu: Diagram height in EMUs (default ~4.2")

        Example:
            SmartArt.add_radial(doc, "Microservices", "API Gateway", [
                "Auth Service",
                "User Service",
                "Payment Service",
                "Notification Service",
            ])
        """
        all_items = [center] + items
        _add_smartart(doc, "radial", all_items, title=title,
                      width_emu=width_emu, height_emu=height_emu)

    @staticmethod
    def finalize(path: str):
        """
        Post-process a saved .docx file to regenerate SmartArt rendering.

        Call this after doc.save() to ensure all SmartArt diagrams render
        correctly on first open. Requires Microsoft Word to be installed.

        Without this step, SmartArt diagrams are embedded as data but may
        appear collapsed until the user clicks on them in Word.

        Args:
            path: Path to the saved .docx file

        Example:
            doc.save("output.docx")
            SmartArt.finalize("output.docx")
        """
        try:
            import win32com.client
        except ImportError:
            import warnings
            warnings.warn(
                "pywin32 not installed. SmartArt diagrams may appear collapsed "
                "until clicked in Word. Install pywin32 and call SmartArt.finalize() "
                "to pre-render diagrams: pip install pywin32"
            )
            return

        import time
        abs_path = os.path.abspath(path)
        word = win32com.client.DispatchEx('Word.Application')
        time.sleep(2)
        word.Visible = False
        try:
            doc = word.Documents.Open(abs_path)
            # Opening and saving forces Word to regenerate all SmartArt drawings
            doc.Save()
            doc.Close(False)
        finally:
            word.Quit()
